from docx import Document
from docx.shared import Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls
from docx.oxml.ns import qn
import json

def load_document(path):
    
    # Create or load an existing document
    doc = Document(path) # Document('existing_doc.docx')
    return doc


def add_images_with_captions(doc, image_list_dict, section_data, max_image_width=3.1, ):
    """
    Create a 2-column table with images and their corresponding captions.
    """
    # Calculate number of rows needed (2 images per row)
    num_rows = (len(image_list_dict) + 1) // 2
    table = doc.add_table(rows=num_rows, cols=2)

    # Track the current row and column index
    current_col = 0
    current_row = 0

    for img_data in image_list_dict:
        # Add image in the current cell
        cell = table.cell(current_row, current_col)
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()

        # Add the image (ensure width is within limit)
        run.add_picture(img_data['new_path'], width=Inches(max_image_width))

        # Add caption below the image

        # Get section numbers:
        section = img_data['section']
        element = img_data['element']
        section_number = section_data[section][element]
        caption_paragraph = cell.add_paragraph(f"Pic {img_data['id']}: {img_data['task']} - ({section_number})")
        caption_paragraph.bold = True

        # Set the cell borders
        set_cell_border(cell, top="single", bottom="single", left="single", right="single")

        # Set padding/margin for the cell
        set_cell_padding(cell, top=100, bottom=100, start=100, end=100)  # Adjust these values for the desired padding

        # Move to the next column
        current_col += 1
        if current_col == 2:  # Once 2 columns are filled, move to the next row
            current_col = 0
            current_row += 1

    # Save the document with the layout
    doc.save('image_grid_with_captions.docx')


def set_cell_border(cell, **kwargs):
    """
    Set borders for a table cell.

    Usage: set_cell_border(cell, top="single", bottom="single", start="single", end="single", size="4", color="000000")
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    for edge in ('top', 'bottom', 'left', 'right'):
        if edge in kwargs:
            # Create a border element for each side
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs.get(edge))
            element.set(qn('w:sz'), kwargs.get('size', '4'))  # Border size (4 = 0.5pt)
            element.set(qn('w:color'), kwargs.get('color', '000000'))  # Border color
            element.set(qn('w:space'), '0')
            tcPr.append(element)


def set_cell_padding(cell, top=0, bottom=0, start=0, end=0):
    """
    Set padding (cell margins) for a table cell.
    """
    tc = cell._element
    tcPr = tc.get_or_add_tcPr()

    cell_margin = OxmlElement('w:tcMar')
    
    # Define margins
    for side in ('top', 'bottom', 'start', 'end'):
        margin = OxmlElement(f'w:{side}')
        margin.set(qn('w:w'), str(locals()[side]))  # Margins in twips (1 pt = 20 twips)
        margin.set(qn('w:type'), 'dxa')
        cell_margin.append(margin)

    tcPr.append(cell_margin)